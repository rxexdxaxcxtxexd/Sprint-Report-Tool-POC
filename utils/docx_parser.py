"""
DOCX parser for Sprint Report Guide documents.

This module provides utilities for parsing Microsoft Word (.docx) files,
specifically designed for extracting Sprint Report Guide templates. Features:
- Extract all text content with preserved structure
- Identify and extract sections by headings
- Validate guide has expected sections
- Handle tables and complex formatting

Example:
    >>> from utils.docx_parser import parse_sprint_guide
    >>> guide_content = parse_sprint_guide("guide.docx")
    >>> sections = extract_sections("guide.docx")
    >>> validation = validate_guide("guide.docx")
"""

import logging
from pathlib import Path
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl


# Module logger
logger = logging.getLogger(__name__)


class DOCXParsingError(Exception):
    """Raised when DOCX parsing fails."""
    pass


# Expected sections in a Sprint Report Guide
EXPECTED_SECTIONS = [
    "Sprint Overview",
    "Completed Work",
    "In Progress",
    "Blockers and Risks",
    "Blocked Items",
    "Metrics",
    "Next Sprint Plan",
    "Sprint Goals"
]


def parse_sprint_guide(docx_path: str) -> str:
    """
    Parse DOCX file and extract all text content.

    Preserves document structure including headings, paragraphs, and tables.
    Ideal for creating a complete text representation of the Sprint Report Guide.

    Args:
        docx_path: Path to DOCX file

    Returns:
        Plain text content with preserved structure (headings, paragraphs, tables)

    Raises:
        FileNotFoundError: If file doesn't exist
        DOCXParsingError: If file is not a valid DOCX or parsing fails

    Example:
        >>> guide_text = parse_sprint_guide("CSG_Sprint_Report_Guide.docx")
        >>> print(guide_text[:100])  # First 100 characters
    """
    docx_file = Path(docx_path)

    # Check file exists
    if not docx_file.exists():
        error_msg = f"DOCX file not found: {docx_path}"
        logger.error(error_msg)
        raise FileNotFoundError(error_msg)

    # Check file extension
    if docx_file.suffix.lower() != ".docx":
        error_msg = f"File is not a DOCX file: {docx_path}"
        logger.error(error_msg)
        raise DOCXParsingError(error_msg)

    try:
        # Load document
        doc = Document(docx_path)
        logger.info(f"Loaded DOCX file: {docx_path}")

        # Extract content
        content_parts = []

        for element in doc.element.body:
            if isinstance(element, CT_P):
                # Paragraph
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()

                if text:
                    # Check if it's a heading
                    is_heading, level = _is_heading(paragraph)

                    if is_heading:
                        # Add heading with markdown-style formatting
                        heading_marker = "#" * level
                        content_parts.append(f"\n{heading_marker} {text}\n")
                    else:
                        # Regular paragraph
                        content_parts.append(text)

            elif isinstance(element, CT_Tbl):
                # Table
                table = Table(element, doc)
                table_text = _extract_table_text(table)
                if table_text:
                    content_parts.append(f"\n{table_text}\n")

        # Join all parts
        full_text = "\n".join(content_parts)

        # Clean up excessive whitespace
        full_text = _clean_whitespace(full_text)

        logger.info(f"Extracted {len(full_text)} characters from {docx_path}")
        return full_text

    except Exception as e:
        error_msg = f"Error parsing DOCX file: {e}"
        logger.error(error_msg, exc_info=True)
        raise DOCXParsingError(error_msg) from e


def extract_sections(docx_path: str) -> Dict[str, str]:
    """
    Parse DOCX and identify sections by headings.

    Groups content under each heading, making it easy to access specific
    sections of the Sprint Report Guide.

    Args:
        docx_path: Path to DOCX file

    Returns:
        Dictionary mapping section names (headings) to their content

    Raises:
        FileNotFoundError: If file doesn't exist
        DOCXParsingError: If file is not a valid DOCX or parsing fails

    Example:
        >>> sections = extract_sections("guide.docx")
        >>> print(sections["Sprint Overview"])
        >>> print(sections.keys())  # All section names
    """
    docx_file = Path(docx_path)

    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        doc = Document(docx_path)
        logger.info(f"Extracting sections from: {docx_path}")

        sections = {}
        current_section = None
        current_content = []

        for element in doc.element.body:
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()

                if not text:
                    continue

                is_heading, level = _is_heading(paragraph)

                if is_heading and level <= 2:
                    # Save previous section
                    if current_section:
                        sections[current_section] = "\n".join(current_content).strip()

                    # Start new section
                    current_section = text
                    current_content = []
                    logger.debug(f"Found section: {current_section}")

                else:
                    # Add content to current section
                    if current_section:
                        current_content.append(text)

            elif isinstance(element, CT_Tbl):
                # Add table content to current section
                if current_section:
                    table = Table(element, doc)
                    table_text = _extract_table_text(table)
                    if table_text:
                        current_content.append(table_text)

        # Save last section
        if current_section:
            sections[current_section] = "\n".join(current_content).strip()

        logger.info(f"Extracted {len(sections)} sections from {docx_path}")
        return sections

    except Exception as e:
        error_msg = f"Error extracting sections: {e}"
        logger.error(error_msg, exc_info=True)
        raise DOCXParsingError(error_msg) from e


def validate_guide(docx_path: str) -> Dict[str, Any]:
    """
    Validate Sprint Report Guide has expected sections.

    Checks if the document contains all expected sections for a complete
    Sprint Report Guide. Useful for verifying guide templates.

    Args:
        docx_path: Path to DOCX file

    Returns:
        Dictionary with validation results:
            - valid: True if all expected sections found
            - sections_found: List of section names found in document
            - missing_sections: List of expected sections not found
            - extra_sections: List of sections found but not expected
            - warnings: List of warning messages
            - paragraph_count: Total number of paragraphs
            - table_count: Total number of tables

    Raises:
        FileNotFoundError: If file doesn't exist
        DOCXParsingError: If file is not a valid DOCX or parsing fails

    Example:
        >>> validation = validate_guide("guide.docx")
        >>> if not validation["valid"]:
        ...     print(f"Missing: {validation['missing_sections']}")
        >>> print(f"Found {len(validation['sections_found'])} sections")
    """
    docx_file = Path(docx_path)

    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        # Extract sections
        sections = extract_sections(docx_path)
        sections_found = list(sections.keys())

        # Load document for statistics
        doc = Document(docx_path)

        # Count paragraphs and tables
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        # Find missing and extra sections
        missing_sections = []
        for expected in EXPECTED_SECTIONS:
            # Check for exact match or partial match
            found = any(
                expected.lower() in section.lower()
                for section in sections_found
            )
            if not found:
                missing_sections.append(expected)

        extra_sections = []
        for found in sections_found:
            # Check if this section matches any expected section
            matches = any(
                expected.lower() in found.lower()
                for expected in EXPECTED_SECTIONS
            )
            if not matches:
                extra_sections.append(found)

        # Generate warnings
        warnings = []

        if not sections_found:
            warnings.append("No sections found - document may not have headings")

        if paragraph_count < 5:
            warnings.append(f"Very short document - only {paragraph_count} paragraphs")

        for section_name, content in sections.items():
            if len(content) < 50:
                warnings.append(
                    f"Section '{section_name}' has minimal content ({len(content)} chars)"
                )

        # Determine if valid
        valid = len(missing_sections) == 0

        result = {
            "valid": valid,
            "sections_found": sections_found,
            "missing_sections": missing_sections,
            "extra_sections": extra_sections,
            "warnings": warnings,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "file_path": str(docx_file.absolute())
        }

        # Log results
        if valid:
            logger.info(f"Validation passed for {docx_path}")
        else:
            logger.warning(
                f"Validation failed for {docx_path}: "
                f"missing {len(missing_sections)} sections"
            )

        return result

    except Exception as e:
        error_msg = f"Error validating guide: {e}"
        logger.error(error_msg, exc_info=True)
        raise DOCXParsingError(error_msg) from e


def get_section(docx_path: str, section_name: str) -> Optional[str]:
    """
    Get content of a specific section from DOCX file.

    Args:
        docx_path: Path to DOCX file
        section_name: Name of section to retrieve (case-insensitive)

    Returns:
        Section content or None if section not found

    Example:
        >>> overview = get_section("guide.docx", "Sprint Overview")
        >>> if overview:
        ...     print(overview)
    """
    sections = extract_sections(docx_path)

    # Try exact match first
    if section_name in sections:
        return sections[section_name]

    # Try case-insensitive match
    for name, content in sections.items():
        if name.lower() == section_name.lower():
            return content

    # Try partial match
    for name, content in sections.items():
        if section_name.lower() in name.lower():
            logger.debug(
                f"Section '{section_name}' matched to '{name}' via partial match"
            )
            return content

    logger.warning(f"Section '{section_name}' not found in {docx_path}")
    return None


def _is_heading(paragraph: Paragraph) -> Tuple[bool, int]:
    """
    Check if paragraph is a heading and return its level.

    Args:
        paragraph: Paragraph object to check

    Returns:
        Tuple of (is_heading: bool, level: int)
        Level is 0 for non-headings, 1-9 for headings

    Example:
        >>> is_heading, level = _is_heading(paragraph)
        >>> if is_heading:
        ...     print(f"Heading level {level}")
    """
    style_name = paragraph.style.name.lower() if paragraph.style else ""

    # Check for heading styles
    if "heading" in style_name:
        # Extract level from style name (e.g., "Heading 1" -> 1)
        try:
            level = int(style_name.split()[-1])
            return True, level
        except (ValueError, IndexError):
            # Default to level 1 if can't parse
            return True, 1

    # Check for title style
    if "title" in style_name:
        return True, 1

    # Check for bold + larger font (heuristic for headings)
    if paragraph.runs:
        first_run = paragraph.runs[0]
        if first_run.bold and first_run.font.size:
            # If significantly larger than normal text (> 12pt)
            point_size = first_run.font.size.pt if first_run.font.size else 0
            if point_size > 12:
                return True, 2

    return False, 0


def _extract_table_text(table: Table) -> str:
    """
    Convert DOCX table to readable plain text.

    Formats table as pipe-separated values for readability.

    Args:
        table: Table object from python-docx

    Returns:
        Formatted table text

    Example:
        >>> table_text = _extract_table_text(table)
        >>> print(table_text)
        # Column1 | Column2 | Column3
        # --------|---------|--------
        # Value1  | Value2  | Value3
    """
    if not table.rows:
        return ""

    lines = []

    # Process each row
    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # Get cell text (may have multiple paragraphs)
            cell_text = " ".join(
                p.text.strip()
                for p in cell.paragraphs
                if p.text.strip()
            )
            cells.append(cell_text)

        # Join cells with pipe separator
        line = " | ".join(cells)
        lines.append(line)

        # Add separator after header row
        if i == 0 and len(table.rows) > 1:
            separator = " | ".join(["-" * max(3, len(cell)) for cell in cells])
            lines.append(separator)

    return "\n".join(lines)


def _clean_whitespace(text: str) -> str:
    """
    Clean up excessive whitespace in text.

    Removes:
    - Multiple consecutive blank lines (keeps max 2)
    - Trailing whitespace on each line
    - Leading/trailing whitespace from entire text

    Args:
        text: Text to clean

    Returns:
        Cleaned text

    Example:
        >>> cleaned = _clean_whitespace("Line 1\\n\\n\\n\\nLine 2  \\n")
        >>> print(cleaned)  # "Line 1\\n\\nLine 2"
    """
    # Split into lines
    lines = text.split("\n")

    # Remove trailing whitespace from each line
    lines = [line.rstrip() for line in lines]

    # Remove excessive blank lines (keep max 2 consecutive)
    cleaned_lines = []
    blank_count = 0

    for line in lines:
        if not line:
            blank_count += 1
            if blank_count <= 2:
                cleaned_lines.append(line)
        else:
            blank_count = 0
            cleaned_lines.append(line)

    # Join and strip
    return "\n".join(cleaned_lines).strip()


def get_document_stats(docx_path: str) -> Dict[str, Any]:
    """
    Get statistics about a DOCX document.

    Args:
        docx_path: Path to DOCX file

    Returns:
        Dictionary with document statistics:
            - paragraph_count: Number of paragraphs
            - table_count: Number of tables
            - heading_count: Number of headings
            - word_count: Approximate word count
            - character_count: Character count

    Example:
        >>> stats = get_document_stats("guide.docx")
        >>> print(f"Document has {stats['word_count']} words")
    """
    docx_file = Path(docx_path)

    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        doc = Document(docx_path)

        # Count elements
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        # Count headings
        heading_count = sum(
            1 for p in doc.paragraphs
            if _is_heading(p)[0]
        )

        # Get text for word/character count
        text = parse_sprint_guide(docx_path)
        word_count = len(text.split())
        character_count = len(text)

        stats = {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "heading_count": heading_count,
            "word_count": word_count,
            "character_count": character_count,
            "file_path": str(docx_file.absolute()),
            "file_size_bytes": docx_file.stat().st_size
        }

        logger.debug(f"Document stats for {docx_path}: {stats}")
        return stats

    except Exception as e:
        error_msg = f"Error getting document stats: {e}"
        logger.error(error_msg, exc_info=True)
        raise DOCXParsingError(error_msg) from e


if __name__ == "__main__":
    # Example usage and testing
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
    )

    import sys

    if len(sys.argv) > 1:
        docx_path = sys.argv[1]

        try:
            print(f"\n=== Parsing {docx_path} ===\n")

            # Get stats
            stats = get_document_stats(docx_path)
            print("Document Statistics:")
            for key, value in stats.items():
                print(f"  {key}: {value}")

            # Extract sections
            print("\n\nExtracted Sections:")
            sections = extract_sections(docx_path)
            for section_name in sections.keys():
                print(f"  - {section_name}")

            # Validate
            print("\n\nValidation Results:")
            validation = validate_guide(docx_path)
            print(f"  Valid: {validation['valid']}")
            if validation['missing_sections']:
                print(f"  Missing: {validation['missing_sections']}")
            if validation['warnings']:
                print(f"  Warnings: {validation['warnings']}")

            # Show first 500 chars of full text
            print("\n\nFirst 500 characters of extracted text:")
            full_text = parse_sprint_guide(docx_path)
            print(full_text[:500])
            print("...")

        except Exception as e:
            print(f"Error: {e}")
            sys.exit(1)
    else:
        print("Usage: python docx_parser.py <path-to-docx-file>")
        print("\nExample:")
        print("  python docx_parser.py guide.docx")
